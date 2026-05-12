from docx import Document as DocxDocument
from pathlib import Path
from app.services.template_analyzer import TemplateAnalyzer
from app.services.template_filler import TemplateFiller


def test_fill_bracket(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("甲方：【XX 公司[姓名]】")
    template_path = str(tmp_path / "template.docx")
    doc.save(template_path)

    fields = TemplateAnalyzer().analyze(template_path)
    assert len(fields) == 1
    field = fields[0]

    filler = TemplateFiller()
    output_path = filler.fill(
        template_path=template_path,
        field_values={field.id: "张三"},
        field_registry={field.id: field},
        output_dir=str(tmp_path),
    )

    result_doc = DocxDocument(output_path)
    text = " ".join(p.text for p in result_doc.paragraphs)
    assert "张三" in text


def test_empty_value_skipped(tmp_path):
    doc = DocxDocument()
    doc.add_paragraph("联系人：________")
    template_path = str(tmp_path / "template.docx")
    doc.save(template_path)

    fields = TemplateAnalyzer().analyze(template_path)
    filler = TemplateFiller()
    output_path = filler.fill(
        template_path=template_path,
        field_values={fields[0].id: ""},
        field_registry={fields[0].id: fields[0]},
        output_dir=str(tmp_path),
    )
    result_doc = DocxDocument(output_path)
    text = " ".join(p.text for p in result_doc.paragraphs)
    assert "________" in text

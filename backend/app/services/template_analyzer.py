"""模板字段分析器 — 识别文档中的待填字段。"""
import re
import uuid
from dataclasses import dataclass
from docx import Document as DocxDocument


@dataclass
class TemplateField:
    id: str
    label: str
    field_type: str  # bracket | blank | table_cell | inline_paren | standalone_blank
    original_text: str
    location: str = ""


# Patterns
BRACKET_PATTERN = re.compile(r'【[^【】]*\[([^\]]+)\][^【】]*】')
BLANK_PATTERN = re.compile(r'([^：:]+)[：:]\s*(_{4,}|—{4,}|-{4,}|\s{4,})')
INLINE_PAREN_PATTERN = re.compile(r'（([^）]{2,30})）')
# Standalone blanks: ______ or ——— in text (no label prefix needed)
STANDALONE_BLANK_PATTERN = re.compile(r'_{4,}|—{4,}|—{4,}|-{4,}')
# Blank with spaces after colon: 标签：      （空格填空）
COLON_SPACE_PATTERN = re.compile(r'([^：:]{2,20})[：:]\s{4,}')
# Date placeholders: 年　月　日 or 年   月   日
DATE_PLACEHOLDER_PATTERN = re.compile(r'(\d{2,4})\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日')


class TemplateAnalyzer:
    def analyze(self, file_path: str) -> list[TemplateField]:
        """分析文档，返回去重后的待填字段列表。"""
        doc = DocxDocument(file_path)
        fields: list[TemplateField] = []
        seen_labels: set[str] = set()

        # Paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            for f in self._extract_from_text(text, f"paragraph_{i}"):
                if f.label not in seen_labels:
                    fields.append(f)
                    seen_labels.add(f.label)

        # Tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    loc = f"table_{t_idx}_row{r_idx}_col{c_idx}"
                    for f in self._extract_from_text(text, loc):
                        if f.label not in seen_labels:
                            fields.append(f)
                            seen_labels.add(f.label)

        # Empty table cells (potential answer cells)
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text and len(row.cells) > 1:
                        # Empty cell in a multi-column table — likely an answer cell
                        # Use header row or first row as label context
                        label = self._guess_cell_label(table, r_idx, c_idx)
                        if label and label not in seen_labels:
                            loc = f"table_{t_idx}_row{r_idx}_col{c_idx}"
                            fields.append(TemplateField(
                                id=str(uuid.uuid4()),
                                label=label,
                                field_type="table_cell",
                                original_text="",
                                location=loc,
                            ))
                            seen_labels.add(label)

        return fields

    def _guess_cell_label(self, table, row_idx: int, col_idx: int) -> str | None:
        """从表头或同行其他列猜测空白单元格的标签。"""
        # Try header row (first row)
        if table.rows:
            header_row = table.rows[0]
            if col_idx < len(header_row.cells):
                header_text = header_row.cells[col_idx].text.strip()
                if header_text:
                    return header_text[:30]
        # Try same row, first non-empty cell
        row = table.rows[row_idx]
        for i, cell in enumerate(row.cells):
            if i != col_idx and cell.text.strip():
                return cell.text.strip()[:30]
        return None

    def _extract_from_text(self, text: str, location: str) -> list[TemplateField]:
        """从单段文本中提取所有字段。"""
        results: list[TemplateField] = []

        # Bracket: 【XX[姓名]】
        for m in BRACKET_PATTERN.finditer(text):
            label = m.group(1).strip()
            if label:
                results.append(TemplateField(
                    id=str(uuid.uuid4()),
                    label=label,
                    field_type="bracket",
                    original_text=m.group(0),
                    location=location,
                ))

        # Blank: 标签：________ or 标签：      (spaces)
        for m in BLANK_PATTERN.finditer(text):
            label = m.group(1).strip()
            if label:
                results.append(TemplateField(
                    id=str(uuid.uuid4()),
                    label=label,
                    field_type="blank",
                    original_text=m.group(0),
                    location=location,
                ))

        # Colon-space: 标签：      (4+ spaces after colon)
        for m in COLON_SPACE_PATTERN.finditer(text):
            label = m.group(1).strip()
            if label and not any(r.label == label for r in results):
                results.append(TemplateField(
                    id=str(uuid.uuid4()),
                    label=label,
                    field_type="blank",
                    original_text=m.group(0),
                    location=location,
                ))

        # Inline paren: （投标人名称）or （    ）
        for m in INLINE_PAREN_PATTERN.finditer(text):
            label = m.group(1).strip()
            if label and not any(r.label == label for r in results):
                # Skip if it's just spaces
                if len(label) >= 2 and not label.isspace():
                    results.append(TemplateField(
                        id=str(uuid.uuid4()),
                        label=label,
                        field_type="inline_paren",
                        original_text=m.group(0),
                        location=location,
                    ))

        # Standalone blanks: ______ (use truncated text as label)
        if not results:
            for m in STANDALONE_BLANK_PATTERN.finditer(text):
                # Use text before the blank as label context
                prefix = text[:m.start()].strip()
                # Remove leading numbers like "1." "2."
                prefix = re.sub(r'^\d+[.、)）]\s*', '', prefix)
                if prefix and len(prefix) >= 2:
                    label = prefix[:40]
                    if not any(r.label == label for r in results):
                        results.append(TemplateField(
                            id=str(uuid.uuid4()),
                            label=label,
                            field_type="standalone_blank",
                            original_text=m.group(0),
                            location=location,
                        ))

        return results

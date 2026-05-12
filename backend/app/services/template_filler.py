"""模板填充器 — 将字段值写入 Word 文档，生成新文件。"""
import re
import os
from pathlib import Path
from docx import Document as DocxDocument
from app.services.template_analyzer import TemplateField


class TemplateFiller:
    def fill(
        self,
        template_path: str,
        field_values: dict[str, str],
        field_registry: dict[str, TemplateField],
        output_dir: str,
    ) -> str:
        """将字段值填入模板，返回新文件路径。不修改原始模板。"""
        doc = DocxDocument(template_path)

        for field_id, value in field_values.items():
            if not value:
                continue
            f = field_registry.get(field_id)
            if not f:
                continue
            if f.field_type == "bracket":
                self._fill_bracket(doc, f, value)
            elif f.field_type == "blank":
                self._fill_blank(doc, f, value)
            elif f.field_type == "inline_paren":
                self._fill_inline_paren(doc, f, value)
            elif f.field_type == "table_cell":
                self._fill_table_cell(doc, f, value)

        os.makedirs(output_dir, exist_ok=True)
        output_path = str(Path(output_dir) / f"filled_{Path(template_path).name}")
        doc.save(output_path)
        return output_path

    def _fill_bracket(self, doc: DocxDocument, f: TemplateField, value: str):
        """替换方括号占位符：【XX[姓名]】 → 【XX张三】"""
        pattern = re.compile(re.escape(f.original_text))
        self._replace_in_paragraphs(doc, pattern, value, mode="bracket")

    def _fill_blank(self, doc: DocxDocument, f: TemplateField, value: str):
        """替换下划线空白：联系人：________ → 联系人：张三"""
        pattern = re.compile(re.escape(f.original_text))
        self._replace_in_paragraphs(doc, pattern, f"{f.label}：{value}", mode="blank")

    def _fill_inline_paren(self, doc: DocxDocument, f: TemplateField, value: str):
        """替换全角括号：（投标人名称） → 张三"""
        pattern = re.compile(re.escape(f.original_text))
        self._replace_in_paragraphs(doc, pattern, value, mode="paren")

    def _fill_table_cell(self, doc: DocxDocument, f: TemplateField, value: str):
        """替换表格单元格中的占位符。"""
        pattern = re.compile(re.escape(f.original_text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, pattern, value)

    def _replace_in_paragraphs(self, doc: DocxDocument, pattern: re.Pattern, replacement: str, mode: str):
        """在所有段落中执行替换（跨 run 处理）。"""
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, pattern, replacement)
        # Also check tables, headers, footers
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, pattern, replacement)

    def _replace_in_paragraph(self, paragraph, pattern: re.Pattern, replacement: str):
        """在单个段落中执行 run 级别替换。"""
        full_text = paragraph.text
        if not pattern.search(full_text):
            return
        # Simple approach: rebuild runs
        new_text = pattern.sub(replacement, full_text)
        if paragraph.runs:
            # Clear all runs except first, set text on first
            for run in paragraph.runs[1:]:
                run.text = ""
            paragraph.runs[0].text = new_text

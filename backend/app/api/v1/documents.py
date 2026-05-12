"""文档 CRUD + AI 填写触发端点。"""
import logging
import uuid
import json
from pathlib import Path

import aiofiles
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.database import SessionLocal
from app.core.deps import get_db
from app.models.document import Document
from app.schemas.document import (
    ConfirmResponse,
    DocumentStatusResponse,
    DocumentUploadResponse,
    EditorTokenResponse,
    FieldUpdateRequest,
)
from app.services.ai_filler import AiFiller
from app.services.onlyoffice_service import OnlyOfficeService
from app.services.template_analyzer import TemplateAnalyzer
from app.services.template_filler import TemplateFiller
from app.services.word_parser import WordParser

router = APIRouter()
logger = logging.getLogger(__name__)
onlyoffice_service = OnlyOfficeService()
ai_filler = AiFiller()

ALLOWED_EXTENSIONS = {".docx", ".doc"}


def _extract_text(file_path: str) -> str:
    """从 docx 文件提取纯文本。"""
    try:
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """上传目标文档（待填写的 Word 文件）。"""
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"不支持的文件类型: {suffix}，请上传 .docx 或 .doc 文件")

    doc_id = str(uuid.uuid4())
    upload_dir = settings.upload_dir / "target"
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / f"{doc_id}{suffix}"

    async with aiofiles.open(file_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    try:
        analyzer = TemplateAnalyzer()
        fields_raw = analyzer.analyze(str(file_path))
        fields = [
            {
                "id": f.id,
                "label": f.label,
                "value": "",
                "status": "empty",
                "field_type": f.field_type,
                "requires_input": False,
                "original_text": f.original_text,
                "location": f.location,
            }
            for f in fields_raw
        ]

        parser = WordParser(str(file_path))
        outline = parser.extract_outline()
        parser.doc.save(str(file_path))

        status = "ready" if fields else "error"
        error_message = None if fields else "未识别到可填写字段"
    except Exception as e:
        logger.error("解析文档失败: %s", e)
        fields = []
        outline = []
        status = "error"
        error_message = str(e)

    doc = Document(
        id=doc_id,
        original_filename=file.filename or "document.docx",
        file_path=str(file_path),
        status=status,
        fields=fields,
        outline=outline,
        references=[],
        error_message=error_message,
    )
    db.add(doc)
    db.commit()

    return DocumentUploadResponse(
        doc_id=doc_id,
        status=status,
        message="文档上传成功，字段解析完成" if fields else "文档上传成功，但未识别到可填写字段",
    )


@router.post("/{doc_id}/references", response_model=DocumentUploadResponse)
async def upload_reference(
    doc_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    """上传参考文档。"""
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "文档不存在")

    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"不支持的文件类型: {suffix}")

    ref_id = str(uuid.uuid4())
    ref_dir = settings.upload_dir / "references" / doc_id
    ref_dir.mkdir(parents=True, exist_ok=True)
    ref_path = ref_dir / f"{ref_id}{suffix}"

    async with aiofiles.open(ref_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    ref_text = _extract_text(str(ref_path))
    refs = list(doc.references or [])
    refs.append({
        "doc_id": ref_id,
        "filename": file.filename,
        "file_path": str(ref_path),
        "text": ref_text[:10000],
    })
    doc.references = refs
    db.commit()

    return DocumentUploadResponse(
        doc_id=ref_id,
        status="ready",
        message="参考文档上传成功",
    )


@router.get("/{doc_id}", response_model=DocumentStatusResponse)
def get_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "文档不存在")
    return DocumentStatusResponse(
        doc_id=doc.id,
        original_filename=doc.original_filename,
        status=doc.status,
        fields=doc.fields,
        outline=doc.outline,
        references=[
            {"doc_id": r["doc_id"], "filename": r["filename"]}
            for r in (doc.references or [])
        ],
        error_message=doc.error_message,
    )


@router.get("/{doc_id}/editor-token", response_model=EditorTokenResponse)
def get_editor_token(doc_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "文档不存在")
    result = onlyoffice_service.generate_editor_config(
        doc_id=doc_id,
        filename=doc.original_filename,
        host_url=settings.host_url,
    )
    doc.onlyoffice_doc_key = result["doc_key"]
    db.commit()
    return EditorTokenResponse(**result)


@router.get("/{doc_id}/raw-file")
def get_raw_file(doc_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc or not Path(doc.file_path).exists():
        raise HTTPException(404, "文件不存在")
    return FileResponse(
        doc.file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.post("/{doc_id}/ai-fill")
def trigger_ai_fill(doc_id: str, resume: bool = False, db: Session = Depends(get_db)):
    """触发 AI 填写，返回 SSE 流。支持断点续传。"""
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "文档不存在")
    if not doc.fields:
        raise HTTPException(400, "文档无可填写字段")

    start_chunk = 0
    partial: dict[str, str] = {}
    if resume and doc.fill_progress:
        start_chunk = doc.fill_progress.get("chunk_index", 0)
        partial = doc.partial_fields or {}
        doc.fill_progress = {**doc.fill_progress, "cancelled": False}
        doc.status = "filling"
        db.commit()
    else:
        doc.fill_progress = None
        doc.partial_fields = None
        doc.status = "filling"
        db.commit()

    fields = list(doc.fields)
    for f in fields:
        if f["id"] in partial:
            f["value"] = partial[f["id"]]

    document_text = _extract_text(doc.file_path)
    reference_text = None
    if doc.references:
        reference_text = "\n\n---\n\n".join(
            r["text"] for r in doc.references if r.get("text")
        )

    def event_stream():
        gen_db = SessionLocal()
        errored = False
        error_message = None
        try:
            for sse_line in ai_filler.fill_stream(
                fields=fields,
                document_text=document_text,
                reference_text=reference_text or None,
                start_chunk=start_chunk,
                partial_result=partial if resume else None,
            ):
                if sse_line.startswith("event: field_filled"):
                    data_line = sse_line.split("\ndata: ", 1)[1].strip()
                    item = json.loads(data_line.split("\n\n")[0])
                    partial[item["id"]] = item["value"]
                elif sse_line.startswith("event: field_requires_input"):
                    data_line = sse_line.split("\ndata: ", 1)[1].strip()
                    item = json.loads(data_line.split("\n\n")[0])
                    partial.setdefault(item["id"], "")
                elif sse_line.startswith("event: progress"):
                    data_line = sse_line.split("\ndata: ", 1)[1].strip()
                    progress = json.loads(data_line.split("\n\n")[0])
                    progress_doc = gen_db.query(Document).filter(Document.id == doc_id).first()
                    if progress_doc:
                        progress_doc.fill_progress = {
                            "chunk_index": progress.get("chunk_index", 0),
                            "total_chunks": progress.get("total_chunks", 0),
                            "cancelled": bool((progress_doc.fill_progress or {}).get("cancelled", False)),
                        }
                        progress_doc.partial_fields = dict(partial)
                        gen_db.commit()
                elif sse_line.startswith("event: error"):
                    errored = True
                    data_line = sse_line.split("\ndata: ", 1)[1].strip()
                    error_message = json.loads(data_line.split("\n\n")[0]).get("message")
                yield sse_line

                _check_doc = gen_db.query(Document).filter(Document.id == doc_id).first()
                if _check_doc and _check_doc.fill_progress and _check_doc.fill_progress.get("cancelled"):
                    fields_to_save = list(_check_doc.fields or [])
                    for field in fields_to_save:
                        field_id = field.get("id")
                        if field_id in partial:
                            field["value"] = partial[field_id]
                            field["status"] = "filled" if partial[field_id] else "empty"
                    _check_doc.fields = fields_to_save
                    _check_doc.status = "paused"
                    _check_doc.partial_fields = dict(partial)
                    gen_db.commit()
                    yield f"event: cancelled\ndata: {json.dumps({'message': 'AI 填写已暂停'})}\n\n"
                    return
        finally:
            _doc = gen_db.query(Document).filter(Document.id == doc_id).first()
            if _doc and _doc.status == "filling":
                fields_to_save = list(_doc.fields or [])
                for field in fields_to_save:
                    field_id = field.get("id")
                    if field_id in partial:
                        field["value"] = partial[field_id]
                        field["status"] = "filled" if partial[field_id] else "empty"
                _doc.fields = fields_to_save
                _doc.partial_fields = dict(partial)
                _doc.status = "error" if errored else "ready"
                _doc.error_message = error_message
                gen_db.commit()
            gen_db.close()

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/{doc_id}/ai-fill-cancel")
def cancel_ai_fill(doc_id: str, db: Session = Depends(get_db)):
    """取消 AI 填写。"""
    doc = db.get(Document, doc_id)
    if not doc:
        raise HTTPException(404, "文档不存在")
    doc.fill_progress = {**(doc.fill_progress or {}), "cancelled": True}
    db.commit()
    return {"success": True, "message": "取消请求已提交，AI 将在当前步骤完成后暂停"}


@router.patch("/{doc_id}/fields/{field_id}")
def update_field(
    doc_id: str,
    field_id: str,
    body: FieldUpdateRequest,
    db: Session = Depends(get_db),
):
    doc = db.get(Document, doc_id)
    if not doc or not doc.fields:
        raise HTTPException(404, "文档或字段不存在")

    fields = list(doc.fields)
    found = False
    for f in fields:
        if f["id"] == field_id:
            f["value"] = body.value
            f["status"] = "filled" if body.value else "empty"
            found = True
            break

    if not found:
        raise HTTPException(404, f"字段 {field_id} 不存在")

    doc.fields = fields
    db.commit()
    return {"success": True, "field_id": field_id, "value": body.value}


@router.post("/{doc_id}/confirm", response_model=ConfirmResponse)
def confirm_fields(doc_id: str, db: Session = Depends(get_db)):
    """将字段值写回 Word 文档，生成可下载版本。"""
    doc = db.get(Document, doc_id)
    if not doc or not doc.fields:
        raise HTTPException(400, "文档或字段不存在")

    analyzer = TemplateAnalyzer()
    fields_raw = analyzer.analyze(doc.file_path)
    field_registry = {f.id: f for f in fields_raw}

    field_values = {
        f["id"]: f["value"]
        for f in doc.fields
        if f.get("value")
    }

    output_dir = str(settings.processed_dir / doc_id)
    filler = TemplateFiller()
    output_path = filler.fill(
        template_path=doc.file_path,
        field_values=field_values,
        field_registry=field_registry,
        output_dir=output_dir,
    )

    doc.status = "filled"
    doc.file_path = output_path
    db.commit()

    download_url = f"{settings.host_url}/api/v1/documents/{doc_id}/download"
    return ConfirmResponse(
        success=True,
        download_url=download_url,
        message="字段已写入文档，可以下载",
    )


@router.get("/{doc_id}/download")
def download_document(doc_id: str, db: Session = Depends(get_db)):
    doc = db.get(Document, doc_id)
    if not doc or not Path(doc.file_path).exists():
        raise HTTPException(404, "文件不存在，请先确认字段")
    return FileResponse(
        doc.file_path,
        filename=f"filled_{doc.original_filename}",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
